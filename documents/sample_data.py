"""Self-contained multilingual demonstration documents and example questions."""

from __future__ import annotations

import json

from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


@dataclass(frozen=True)
class SampleTable:
    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]


@dataclass(frozen=True)
class SampleSection:
    heading: str
    paragraphs: tuple[str, ...] = ()
    table: SampleTable | None = None


@dataclass(frozen=True)
class SampleDocument:
    title: str
    filename: str
    heading: str
    sections: tuple[SampleSection, ...]


SAMPLE_DOCUMENTS = (
    SampleDocument(
        title="Neural Radiance Fields",
        filename="neural_radiance_fields.docx",
        heading="Neural Radiance Fields",
        sections=(
            SampleSection(
                heading="Overview",
                paragraphs=(
                    "Neural Radiance Fields (NeRF) use neural networks to represent "
                    "a three-dimensional scene and synthesize new camera views.",
                ),
            ),
            SampleSection(
                heading="Volumetric rendering",
                paragraphs=(
                    "NeRF casts a ray from each camera pixel, samples multiple "
                    "points along the ray, and combines color and density values "
                    "using volumetric rendering.",
                ),
            ),
        ),
    ),
    SampleDocument(
        title="راهنمای استفاده از ownCloud",
        filename="owncloud_user_guide_fa.docx",
        heading="راهنمای جامع سامانه ownCloud",
        sections=(
            SampleSection(
                heading="ورود و امنیت حساب",
                paragraphs=(
                    "کاربر پس از ورود به حساب خود باید تأیید دومرحله‌ای یا MFA را "
                    "از بخش Settings > Security فعال کند.",
                    "پیوند اشتراک‌گذاری باید گذرواژه داشته باشد و حداکثر پس از "
                    "هفت روز منقضی شود.",
                ),
            ),
            SampleSection(
                heading="بارگذاری و همگام‌سازی",
                paragraphs=(
                    "فایل‌های DOCX، TXT و PDF برای پرسش‌وپاسخ پشتیبانی می‌شوند. "
                    "هر سند پس از بارگذاری به بخش‌های قابل جست‌وجو تقسیم می‌شود.",
                    "برنامه دسکتاپ ownCloud تغییرات پوشه انتخاب‌شده را به‌صورت "
                    "خودکار با فضای ابری همگام‌سازی می‌کند.",
                ),
            ),
            SampleSection(
                heading="پلن‌ها و فضای ذخیره‌سازی",
                table=SampleTable(
                    headers=("پلن", "فضای ذخیره", "مبلغ ماهانه", "کد سرویس"),
                    rows=(
                        ("شروع", "۲۰ گیگابایت", "۱۲۰٬۰۰۰ تومان", "OC-START-020"),
                        ("حرفه‌ای Pro", "۵۰۰ گیگابایت", "۴۹۰٬۰۰۰ تومان", "OC-PRO-500"),
                        ("سازمانی Business", "۲ ترابایت", "۱٬۴۹۰٬۰۰۰ تومان", "OC-BIZ-2000"),
                    ),
                ),
            ),
            SampleSection(
                heading="نسخه‌های قبلی و بازیابی",
                paragraphs=(
                    "نسخه‌های قبلی فایل‌ها تا ۳۰ روز نگهداری می‌شوند و از زبانه "
                    "Versions قابل بازیابی هستند.",
                    "فایل حذف‌شده تا ۱۴ روز در Deleted files باقی می‌ماند.",
                ),
            ),
        ),
    ),
    SampleDocument(
        title="شرایط پشتیبانی و قیمت‌گذاری",
        filename="support_and_pricing_fa.docx",
        heading="شرایط پشتیبانی و توافق‌نامه خدمات",
        sections=(
            SampleSection(
                heading="سطوح پشتیبانی و SLA",
                paragraphs=(
                    "زمان پاسخ اولیه از لحظه ثبت درخواست در مرکز پشتیبانی "
                    "محاسبه می‌شود.",
                ),
                table=SampleTable(
                    headers=("سطح", "پاسخ اولیه", "ساعات پوشش", "شناسه SLA"),
                    rows=(
                        ("Standard", "۸ ساعت کاری", "شنبه تا چهارشنبه", "SLA-STD-008"),
                        ("Professional", "۲ ساعت", "هر روز ۸ تا ۲۰", "SLA-PRO-002"),
                        ("Enterprise", "۳۰ دقیقه", "۲۴ ساعته، ۷ روز هفته", "SLA-ENT-247"),
                    ),
                ),
            ),
            SampleSection(
                heading="پرداخت و صورت‌حساب",
                paragraphs=(
                    "فاکتور نمونه با شناسه INV-2026-0456 مربوط به پلن حرفه‌ای "
                    "است و مبلغ آن ۴۹۰٬۰۰۰ تومان است.",
                    "مهلت پرداخت هر فاکتور هفت روز پس از تاریخ صدور است.",
                ),
            ),
            SampleSection(
                heading="پشتیبان‌گیری و بازیابی اطلاعات",
                paragraphs=(
                    "نسخه پشتیبان روزانه ساعت ۰۲:۰۰ ایجاد می‌شود و به مدت "
                    "۳۰ روز نگهداری خواهد شد.",
                    "درخواست بازیابی داده با کد REF-RESTORE-302 ثبت می‌شود.",
                ),
            ),
            SampleSection(
                heading="لغو اشتراک",
                paragraphs=(
                    "کاربر می‌تواند تا ۴۸ ساعت پیش از تمدید اشتراک درخواست "
                    "لغو را ثبت کند.",
                ),
            ),
        ),
    ),
)


SAMPLE_QUESTIONS = (
    {
        "question": "مبلغ ماهانه پلن حرفه‌ای Pro چقدر است؟",
        "document": "owncloud_user_guide_fa.docx",
        "expected_contains": "۴۹۰٬۰۰۰ تومان",
        "section": "پلن‌ها و فضای ذخیره‌سازی",
    },
    {
        "question": "کد سرویس پلن حرفه‌ای چیست؟",
        "document": "owncloud_user_guide_fa.docx",
        "expected_contains": "OC-PRO-500",
        "section": "پلن‌ها و فضای ذخیره‌سازی",
    },
    {
        "question": "تأیید دومرحله‌ای MFA از کدام بخش فعال می‌شود؟",
        "document": "owncloud_user_guide_fa.docx",
        "expected_contains": "Settings > Security",
        "section": "ورود و امنیت حساب",
    },
    {
        "question": "زمان پاسخ اولیه پشتیبانی Enterprise چقدر است؟",
        "document": "support_and_pricing_fa.docx",
        "expected_contains": "۳۰ دقیقه",
        "section": "سطوح پشتیبانی و SLA",
    },
    {
        "question": "شناسه SLA-ENT-247 مربوط به چه سطحی است؟",
        "document": "support_and_pricing_fa.docx",
        "expected_contains": "Enterprise",
        "section": "سطوح پشتیبانی و SLA",
    },
    {
        "question": "فاکتور INV-2026-0456 چه مبلغی دارد؟",
        "document": "support_and_pricing_fa.docx",
        "expected_contains": "۴۹۰٬۰۰۰ تومان",
        "section": "پرداخت و صورت‌حساب",
    },
    {
        "question": "نسخه پشتیبان اطلاعات چند روز نگهداری می‌شود؟",
        "document": "support_and_pricing_fa.docx",
        "expected_contains": "۳۰ روز",
        "section": "پشتیبان‌گیری و بازیابی اطلاعات",
    },
    {
        "question": "What rendering technique does NeRF use?",
        "document": "neural_radiance_fields.docx",
        "expected_contains": "volumetric rendering",
        "section": "Volumetric rendering",
    },
    {
        "question": "پایتخت چین چیست؟",
        "document": None,
        "expected_behavior": "insufficient_context",
        "description": "Unrelated questions must not be answered from these documents.",
    },
)


def _is_persian(text: str) -> bool:
    return any("\u0600" <= character <= "\u06ff" for character in text)


def _set_paragraph_direction(paragraph, text: str) -> None:
    if not _is_persian(text):
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    properties = paragraph._p.get_or_add_pPr()

    if properties.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi") is None:
        properties.append(OxmlElement("w:bidi"))


def build_sample_document(specification: SampleDocument, target: Path) -> Path:
    document = DocxDocument()
    document.core_properties.title = specification.title

    heading = document.add_heading(specification.heading, level=1)
    _set_paragraph_direction(heading, specification.heading)

    for section in specification.sections:
        heading = document.add_heading(section.heading, level=2)
        _set_paragraph_direction(heading, section.heading)

        for text in section.paragraphs:
            paragraph = document.add_paragraph(text)
            _set_paragraph_direction(paragraph, text)

        if section.table is None:
            continue

        table = document.add_table(rows=1, cols=len(section.table.headers))
        table.style = "Light Shading Accent 1"

        for cell, text in zip(table.rows[0].cells, section.table.headers, strict=True):
            cell.text = text
            _set_paragraph_direction(cell.paragraphs[0], text)

        for values in section.table.rows:
            cells = table.add_row().cells

            for cell, text in zip(cells, values, strict=True):
                cell.text = text
                _set_paragraph_direction(cell.paragraphs[0], text)

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target


def ensure_sample_files(directory: Path, *, overwrite: bool = False) -> list[Path]:
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    created: list[Path] = []

    for specification in SAMPLE_DOCUMENTS:
        target = directory / specification.filename

        if overwrite or not target.exists():
            build_sample_document(specification, target)

        created.append(target)

    questions_path = directory / "sample_questions.json"

    if overwrite or not questions_path.exists():
        questions_path.write_text(
            json.dumps(SAMPLE_QUESTIONS, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    return created

