from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.core.management.base import BaseCommand, CommandError
from docx import Document as DocxDocument

from documents.models import Document


def write_sample_docx(source_txt: Path, target: Path) -> None:
    document = DocxDocument()
    document.add_heading("Neural Radiance Fields", level=1)
    for paragraph in source_txt.read_text(encoding="utf-8").split("\n\n"):
        text = paragraph.strip()
        if text:
            document.add_paragraph(text)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


class Command(BaseCommand):
    help = "Load the bundled sample .docx into Django."

    def handle(self, *args, **options):
        sample_dir = settings.SAMPLE_DATA_DIR
        sample_dir.mkdir(parents=True, exist_ok=True)
        source_txt = settings.BASE_DIR / "txt.txt"
        sample_docx = sample_dir / "neural_radiance_fields.docx"
        if not sample_docx.exists():
            if not source_txt.exists():
                raise CommandError(f"Missing sample source: {source_txt}")
            write_sample_docx(source_txt, sample_docx)

        if Document.objects.filter(title="Neural Radiance Fields").exists():
            self.stdout.write("Sample document already exists.")
            return

        with sample_docx.open("rb") as handle:
            document = Document(title="Neural Radiance Fields")
            document.file.save(sample_docx.name, File(handle), save=True)
        self.stdout.write(self.style.SUCCESS(f"Loaded sample document id={document.pk}."))
