import json
import tempfile
import unittest

from pathlib import Path

from docx import Document as DocxDocument

from documents.sample_data import (
    SAMPLE_DOCUMENTS,
    SAMPLE_QUESTIONS,
    _set_paragraph_direction,
    build_sample_document,
    ensure_sample_files,
)
from rag.langchain_rag import split_document_text
from rag.loaders import load_document


class SampleDataTests(unittest.TestCase):
    def test_generator_creates_three_docx_files_and_question_dataset(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            generated = ensure_sample_files(root)

            self.assertEqual(len(generated), 3)
            self.assertTrue(all(path.exists() for path in generated))
            self.assertTrue((root / "sample_questions.json").exists())

    def test_owncloud_guide_contains_searchable_persian_pricing_table(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            ensure_sample_files(root)
            document = DocxDocument(root / "owncloud_user_guide_fa.docx")
            extracted = load_document(root / "owncloud_user_guide_fa.docx")

            self.assertEqual(len(document.tables), 1)
            self.assertIn("پلن: حرفه‌ای Pro", extracted)
            self.assertIn("مبلغ ماهانه: ۴۹۰٬۰۰۰ تومان", extracted)
            self.assertIn("کد سرویس: OC-PRO-500", extracted)

    def test_support_document_contains_sla_table_and_invoice_identifier(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            ensure_sample_files(root)
            extracted = load_document(root / "support_and_pricing_fa.docx")

            self.assertIn("سطح: Enterprise", extracted)
            self.assertIn("پاسخ اولیه: ۳۰ دقیقه", extracted)
            self.assertIn("SLA-ENT-247", extracted)
            self.assertIn("INV-2026-0456", extracted)

    def test_generated_headings_become_section_aware_chunks(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            ensure_sample_files(root)
            text = load_document(root / "support_and_pricing_fa.docx")
            chunks = split_document_text(text, source="document:sample")
            invoice = next(
                chunk for chunk in chunks if "INV-2026-0456" in chunk.page_content
            )

            self.assertEqual(
                invoice.metadata["section_path"],
                "شرایط پشتیبانی و توافق‌نامه خدمات > پرداخت و صورت‌حساب".replace(
                    "\u200c", ""
                ),
            )

    def test_question_dataset_uses_valid_utf8_without_escaped_persian(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            ensure_sample_files(root)
            raw = (root / "sample_questions.json").read_text(encoding="utf-8")
            questions = json.loads(raw)

            self.assertEqual(len(questions), len(SAMPLE_QUESTIONS))
            self.assertIn("مبلغ ماهانه پلن حرفه‌ای", raw)
            self.assertNotIn("\\u06", raw)

    def test_question_expected_answers_exist_in_the_matching_document(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            ensure_sample_files(root)

            for question in SAMPLE_QUESTIONS:
                if question.get("document") is None:
                    continue

                with self.subTest(question=question["question"]):
                    content = load_document(root / question["document"])
                    self.assertIn(question["expected_contains"], content)

    def test_dataset_contains_unrelated_question_for_grounding_evaluation(self) -> None:
        unrelated = [
            question
            for question in SAMPLE_QUESTIONS
            if question.get("expected_behavior") == "insufficient_context"
        ]

        self.assertEqual(len(unrelated), 1)
        self.assertIsNone(unrelated[0]["document"])

    def test_generation_is_idempotent_when_files_already_exist(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            generated = ensure_sample_files(root)
            before = {path.name: path.read_bytes() for path in generated}

            ensure_sample_files(root)

            self.assertEqual(before, {path.name: path.read_bytes() for path in generated})

    def test_overwrite_recreates_a_corrupted_document(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            generated = ensure_sample_files(root)
            generated[1].write_bytes(b"invalid docx")

            ensure_sample_files(root, overwrite=True)

            self.assertIn("ownCloud", load_document(generated[1]))

    def test_document_generation_does_not_require_an_external_text_file(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "standalone.docx"
            build_sample_document(SAMPLE_DOCUMENTS[0], target)

            self.assertIn("volumetric rendering", load_document(target))

    def test_persian_direction_marker_is_not_duplicated(self) -> None:
        document = DocxDocument()
        paragraph = document.add_paragraph("متن فارسی")

        _set_paragraph_direction(paragraph, paragraph.text)
        _set_paragraph_direction(paragraph, paragraph.text)

        markers = paragraph._p.get_or_add_pPr().findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi"
        )
        self.assertEqual(len(markers), 1)
