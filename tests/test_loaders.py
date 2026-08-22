import tempfile
import unittest
from pathlib import Path

from rag.langchain_rag import (
    chunk_ids_for_source,
    dedupe_documents,
    langchain_database_url,
)
from rag.llm import clean_llm_answer
from rag.loaders import load_document


class LoaderTests(unittest.TestCase):
    def test_chunk_ids_for_source(self) -> None:
        self.assertEqual(
            chunk_ids_for_source("document:3", 2),
            ["document:3:0", "document:3:1"],
        )

    def test_langchain_database_url_adds_psycopg(self) -> None:
        self.assertEqual(
            langchain_database_url(
                "postgresql://postgres:postgres@localhost:5432/roshan_rag"
            ),
            "postgresql+psycopg://postgres:postgres@localhost:5432/roshan_rag",
        )

    def test_load_document_reads_txt(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "note.txt"
            path.write_text("hello rag", encoding="utf-8")
            self.assertEqual(load_document(path), "hello rag")

    def test_clean_llm_answer_drops_thinking(self) -> None:
        raw = (
            "Here's a thinking process:\n"
            "1. Analyze User Input\n"
            "NeRF is a model.\n"
            "بر اساس بخش ۰، نرف یک روش گرافیکی سه‌بعدی است."
        )
        cleaned = clean_llm_answer(raw)
        self.assertIn("نرف", cleaned)
        self.assertNotIn("thinking process", cleaned.lower())

    def test_dedupe_documents_keeps_first_copy(self) -> None:
        from langchain_core.documents import Document

        documents = [
            Document(page_content="same text", metadata={"source": "./txt.txt"}),
            Document(page_content="same text", metadata={"source": "document:1"}),
            Document(page_content="other", metadata={"source": "document:1"}),
        ]
        unique = dedupe_documents(documents)
        self.assertEqual(len(unique), 2)
        self.assertEqual(unique[0].metadata["source"], "./txt.txt")

    def test_load_docx_preserves_headings_tables_and_document_order(
        self,
    ) -> None:
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as directory:
            path = (
                Path(
                    directory
                )
                / "structured.docx"
            )

            document = DocxDocument()

            document.add_heading(
                "شرایط قرارداد",
                level=1,
            )

            document.add_paragraph(
                "اطلاعات پیش از جدول."
            )

            table = document.add_table(
                rows=2,
                cols=2,
            )

            table.cell(
                0,
                0,
            ).text = "محصول"

            table.cell(
                0,
                1,
            ).text = "قیمت"

            table.cell(
                1,
                0,
            ).text = "لپ‌تاپ"

            table.cell(
                1,
                1,
            ).text = "۵۰ میلیون"

            document.add_paragraph(
                "اطلاعات پس از جدول."
            )

            document.save(
                path
            )

            extracted = load_document(
                path
            )

            self.assertIn(
                "# شرایط قرارداد",
                extracted,
            )

            self.assertIn(
                "محصول | قیمت",
                extracted,
            )

            self.assertIn(
                "محصول: لپ‌تاپ | قیمت: ۵۰ میلیون",
                extracted,
            )

            self.assertLess(
                extracted.index(
                    "اطلاعات پیش از جدول."
                ),
                extracted.index(
                    "محصول | قیمت"
                ),
            )

            self.assertLess(
                extracted.index(
                    "قیمت: ۵۰ میلیون"
                ),
                extracted.index(
                    "اطلاعات پس از جدول."
                ),
            )

    def test_load_docx_preserves_multiple_heading_levels(
        self,
    ) -> None:
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as directory:
            path = (
                Path(
                    directory
                )
                / "headings.docx"
            )

            document = DocxDocument()

            document.add_heading(
                "راهنمای خدمات",
                level=1,
            )

            document.add_heading(
                "شرایط پرداخت",
                level=2,
            )

            document.add_heading(
                "جریمه دیرکرد",
                level=3,
            )

            document.add_paragraph(
                "مبلغ جریمه پنج درصد است."
            )

            document.save(
                path
            )

            extracted = load_document(
                path
            )

            self.assertIn(
                "# راهنمای خدمات",
                extracted,
            )

            self.assertIn(
                "## شرایط پرداخت",
                extracted,
            )

            self.assertIn(
                "### جریمه دیرکرد",
                extracted,
            )

    def test_load_docx_extracts_nested_tables(
        self,
    ) -> None:
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as directory:
            path = (
                Path(
                    directory
                )
                / "nested.docx"
            )

            document = DocxDocument()

            outer_table = document.add_table(
                rows=1,
                cols=1,
            )

            outer_cell = outer_table.cell(
                0,
                0,
            )

            outer_cell.text = (
                "جزئیات سفارش"
            )

            inner_table = outer_cell.add_table(
                rows=2,
                cols=2,
            )

            inner_table.cell(
                0,
                0,
            ).text = "کد"

            inner_table.cell(
                0,
                1,
            ).text = "وضعیت"

            inner_table.cell(
                1,
                0,
            ).text = "ZX-9182"

            inner_table.cell(
                1,
                1,
            ).text = "تأیید شده"

            document.save(
                path
            )

            extracted = load_document(
                path
            )

            self.assertIn(
                "جزئیات سفارش",
                extracted,
            )

            self.assertIn(
                "کد: ZX-9182",
                extracted,
            )

            self.assertIn(
                "وضعیت: تأیید شده",
                extracted,
            )

    def test_load_docx_does_not_duplicate_merged_cells(
        self,
    ) -> None:
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as directory:
            path = (
                Path(
                    directory
                )
                / "merged.docx"
            )

            document = DocxDocument()

            table = document.add_table(
                rows=1,
                cols=2,
            )

            merged_cell = table.cell(
                0,
                0,
            ).merge(
                table.cell(
                    0,
                    1,
                )
            )

            merged_cell.text = (
                "اطلاعات یکپارچه"
            )

            document.save(
                path
            )

            extracted = load_document(
                path
            )

            self.assertEqual(
                extracted.count(
                    "اطلاعات یکپارچه"
                ),
                1,
            )

    def test_load_text_file_removes_utf8_byte_order_mark(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = (
                Path(
                    directory
                )
                / "bom.txt"
            )

            path.write_text(
                "متن فارسی",
                encoding="utf-8-sig",
            )

            self.assertEqual(
                load_document(
                    path
                ),
                "متن فارسی",
            )

    def test_load_pdf_preserves_top_to_bottom_reading_order(
        self,
    ) -> None:
        import fitz

        with tempfile.TemporaryDirectory() as directory:
            path = (
                Path(
                    directory
                )
                / "ordered.pdf"
            )

            with fitz.open() as document:
                page = document.new_page()

                page.insert_text(
                    (
                        72,
                        180,
                    ),
                    "Second paragraph",
                )

                page.insert_text(
                    (
                        72,
                        80,
                    ),
                    "First paragraph",
                )

                document.save(
                    path
                )

            extracted = load_document(
                path
            )

            self.assertLess(
                extracted.index(
                    "First paragraph"
                ),
                extracted.index(
                    "Second paragraph"
                ),
            )

    def test_load_document_rejects_unsupported_extensions(
        self,
    ) -> None:
        with self.assertRaises(
            ValueError
        ):
            load_document(
                "unsupported.csv"
            )