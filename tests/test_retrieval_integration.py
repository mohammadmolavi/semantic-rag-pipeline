"""Exercise real retrieval components without a database or model download."""

from __future__ import annotations

import os
import re
import tempfile
import unittest

from pathlib import Path
from types import ModuleType
from typing import Any
from unittest.mock import patch

from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.vectorstores import InMemoryVectorStore
from rank_bm25 import BM25Okapi

from rag.hybrid import HybridRetriever
from rag.langchain_rag import split_document_text
from rag.loaders import load_document
from rag.reranking import (
    DEFAULT_RERANKER_MODEL,
    CrossEncoderReranker,
    _load_cross_encoder,
)
from rag.retrieval import tokenize
from rag.services import serialize_sources


CONCEPT_GROUPS = (
    frozenset({"automobile", "car", "vehicle", "خودرو", "ماشین"}),
    frozenset({"invoice", "billing", "payment", "فاکتور", "پرداخت", "مبلغ"}),
    frozenset({"contract", "agreement", "قرارداد", "توافق"}),
    frozenset({"cloud", "storage", "backup", "ابری", "ذخیره", "پشتیبان"}),
    frozenset({"security", "password", "login", "امنیت", "رمز", "ورود"}),
    frozenset({"delivery", "shipment", "order", "ارسال", "سفارش"}),
    frozenset({"penalty", "fine", "جریمه", "دیرکرد"}),
)

REFERENCE_RE = re.compile(r"[A-Za-z]{2,}-\d{2,}")


def make_document(
    text: str,
    source: str,
    chunk_index: int = 0,
    **metadata: object,
) -> Document:
    return Document(
        page_content=text,
        metadata={
            "source": source,
            "chunk_index": chunk_index,
            **metadata,
        },
    )


def document_sources(documents: list[Document]) -> list[str]:
    return [str(document.metadata["source"]) for document in documents]


class DeterministicSemanticEmbeddings(Embeddings):
    """Map synonyms to shared dimensions for reproducible cosine search."""

    @staticmethod
    def _embed(text: str) -> list[float]:
        values = [0.001, *([0.0] * len(CONCEPT_GROUPS))]

        for token in tokenize(text):
            for index, concept in enumerate(CONCEPT_GROUPS, start=1):
                if token in concept:
                    values[index] += 1.0

        return values

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._embed(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed(text)


class MetadataFilteredVectorStore(InMemoryVectorStore):
    """Adapt PostgreSQL-style metadata filters to LangChain's memory store."""

    def __init__(self, documents: list[Document]) -> None:
        super().__init__(DeterministicSemanticEmbeddings())
        self.last_search: dict[str, Any] = {}

        if documents:
            self.add_documents(documents)

    def similarity_search(
        self,
        query: str,
        k: int = 4,
        **kwargs: Any,
    ) -> list[Document]:
        metadata_filter = kwargs.pop("filter", None)
        self.last_search = {
            "query": query,
            "k": k,
            "filter": metadata_filter,
        }

        if isinstance(metadata_filter, dict):

            def matches(document: Document) -> bool:
                return all(
                    document.metadata.get(key) == value
                    for key, value in metadata_filter.items()
                )

            kwargs["filter"] = matches

        elif metadata_filter is not None:
            kwargs["filter"] = metadata_filter

        return super().similarity_search(query, k=k, **kwargs)


class DeterministicCrossEncoder:
    """Score actual question/document pairs without downloading a model."""

    def __init__(
        self,
        *,
        scores: list[object] | None = None,
        error: Exception | None = None,
    ) -> None:
        self.scores = scores
        self.error = error
        self.calls: list[dict[str, Any]] = []

    def predict(
        self,
        pairs: list[tuple[str, str]],
        *,
        batch_size: int,
        show_progress_bar: bool,
    ) -> list[object]:
        self.calls.append(
            {
                "pairs": pairs,
                "batch_size": batch_size,
                "show_progress_bar": show_progress_bar,
            }
        )

        if self.error is not None:
            raise self.error

        if self.scores is not None:
            return self.scores

        results = []

        for question, content in pairs:
            question_tokens = set(tokenize(question))
            content_tokens = set(tokenize(content))
            score = float(len(question_tokens & content_tokens))

            for reference in REFERENCE_RE.findall(question):
                if reference.casefold() in content.casefold():
                    score += 10.0

            results.append(score)

        return results


class RealBM25Tests(unittest.TestCase):
    def setUp(self) -> None:
        self.documents = [
            make_document("Invoice alpha alpha alpha payment.", "document:1"),
            make_document("Invoice alpha payment.", "document:2"),
            make_document("Cloud backup storage guide.", "document:3"),
            make_document("Password security login instructions.", "document:4"),
            make_document("Vehicle delivery information.", "document:5"),
        ]

    def build_retriever(self, **settings: Any) -> HybridRetriever:
        return HybridRetriever(
            MetadataFilteredVectorStore([]),
            lexical_documents=self.documents,
            **settings,
        )

    def test_scores_match_real_rank_bm25_implementation(self) -> None:
        retriever = self.build_retriever()
        results = retriever._bm25_search("alpha")
        corpus = [tokenize(document.page_content) for document in self.documents]
        expected_scores = BM25Okapi(corpus).get_scores(tokenize("alpha"))
        scores_by_source = {
            document.metadata["source"]: score
            for document, score in results
        }

        self.assertAlmostEqual(scores_by_source["document:1"], expected_scores[0])
        self.assertAlmostEqual(scores_by_source["document:2"], expected_scores[1])
        self.assertGreater(scores_by_source["document:1"], scores_by_source["document:2"])

    def test_repeated_query_term_increases_relevance(self) -> None:
        results = self.build_retriever()._bm25_search("alpha")

        self.assertEqual(
            [document.metadata["source"] for document, _ in results],
            ["document:1", "document:2"],
        )

    def test_unmatched_documents_are_excluded(self) -> None:
        self.assertEqual(self.build_retriever()._bm25_search("nonexistent"), [])

    def test_empty_query_does_not_build_a_ranking(self) -> None:
        self.assertEqual(self.build_retriever()._bm25_search("  !!  "), [])

    def test_lexical_candidate_limit_is_applied_after_sorting(self) -> None:
        results = self.build_retriever(lexical_k=1)._bm25_search("alpha")

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0][0].metadata["source"], "document:1")

    def test_source_filter_is_applied_before_bm25_indexing(self) -> None:
        results = self.build_retriever(source="document:2")._bm25_search("alpha")

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0][0].metadata["source"], "document:2")

    def test_blank_documents_are_ignored_without_losing_metadata(self) -> None:
        valid = make_document("Unique reference ZX-9182.", "document:valid")
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            lexical_documents=[make_document("  !!  ", "document:blank"), valid],
        )

        results = retriever._bm25_search("ZX-9182")

        self.assertEqual(len(results), 1)
        self.assertIs(results[0][0], valid)

    def test_single_document_is_kept_even_when_bm25_score_is_negative(self) -> None:
        document = make_document("Contract ZX-9182.", "document:only")
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            lexical_documents=[document],
        )

        results = retriever._bm25_search("ZX-9182")

        self.assertEqual(len(results), 1)
        self.assertIs(results[0][0], document)
        self.assertLess(results[0][1], 0)

    def test_persian_arabic_digits_and_half_spaces_share_the_same_index(self) -> None:
        document = make_document(
            "شماره قرارداد ٤٥٦ ثبت می\u200cشود.",
            "document:persian",
        )
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            lexical_documents=[document],
        )

        results = retriever._bm25_search("قرارداد ۴۵۶ می شود")

        self.assertEqual(len(results), 1)
        self.assertIs(results[0][0], document)


class WeightedRRFIntegrationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.retriever = HybridRetriever(MetadataFilteredVectorStore([]))

    def test_scores_follow_the_weighted_reciprocal_rank_formula(self) -> None:
        semantic = make_document("Semantic document.", "document:semantic")
        shared = make_document("Shared document.", "document:shared")
        lexical = make_document("Lexical document.", "document:lexical")

        ranked = self.retriever._reciprocal_rank_fusion(
            [semantic, shared],
            [(lexical, 900.0), (shared, -100.0)],
        )
        scores = {document.metadata["source"]: score for document, score in ranked}

        self.assertAlmostEqual(scores["document:semantic"], 0.6 / 61)
        self.assertAlmostEqual(scores["document:lexical"], 0.4 / 61)
        self.assertAlmostEqual(scores["document:shared"], 0.6 / 62 + 0.4 / 62)
        self.assertEqual(ranked[0][0].metadata["source"], "document:shared")

    def test_bm25_score_magnitude_does_not_override_rank_position(self) -> None:
        first = make_document("First lexical candidate.", "document:first")
        second = make_document("Second lexical candidate.", "document:second")

        results = self.retriever._reciprocal_rank_fusion(
            [],
            [(first, -500.0), (second, 1_000_000.0)],
        )

        self.assertEqual(
            [document.metadata["source"] for document, _ in results],
            ["document:first", "document:second"],
        )

    def test_duplicate_across_paths_is_merged_by_source_and_chunk(self) -> None:
        vector_copy = make_document("Shared content.", "document:7", 2)
        lexical_copy = make_document("Shared content.", "document:7", 2)

        ranked = self.retriever._reciprocal_rank_fusion(
            [vector_copy],
            [(lexical_copy, 3.0)],
        )

        self.assertEqual(len(ranked), 1)
        self.assertIs(ranked[0][0], vector_copy)
        self.assertAlmostEqual(ranked[0][1], 1 / 61)

    def test_same_content_from_different_sources_is_not_merged(self) -> None:
        first = make_document("Identical content.", "document:one")
        second = make_document("Identical content.", "document:two")

        ranked = self.retriever._reciprocal_rank_fusion([first], [(second, 2.0)])

        self.assertEqual(len(ranked), 2)

    def test_distinct_chunks_from_the_same_source_remain_independent(self) -> None:
        first = make_document("First section.", "document:one", 0)
        second = make_document("Second section.", "document:one", 1)

        ranked = self.retriever._reciprocal_rank_fusion([first], [(second, 2.0)])

        self.assertEqual(
            {document.metadata["chunk_index"] for document, _ in ranked},
            {0, 1},
        )

    def test_duplicate_inside_one_ranking_does_not_inflate_its_score(self) -> None:
        document = make_document("Repeated semantic result.", "document:duplicate")

        ranked = self.retriever._reciprocal_rank_fusion(
            [document, document, document],
            [],
        )

        self.assertEqual(len(ranked), 1)
        self.assertAlmostEqual(ranked[0][1], 0.6 / 61)

    def test_tied_scores_preserve_semantic_insertion_order(self) -> None:
        first = make_document("Vector candidate.", "document:vector")
        second = make_document("Lexical candidate.", "document:lexical")
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            vector_weight=0.5,
            lexical_weight=0.5,
        )

        ranked = retriever._reciprocal_rank_fusion([first], [(second, 100.0)])

        self.assertEqual(
            [document.metadata["source"] for document, _ in ranked],
            ["document:vector", "document:lexical"],
        )

    def test_zero_weight_removes_that_entire_retrieval_path(self) -> None:
        semantic = make_document("Semantic candidate.", "document:semantic")
        lexical = make_document("Lexical candidate.", "document:lexical")
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            vector_weight=0,
            lexical_weight=1,
        )

        ranked = retriever._reciprocal_rank_fusion([semantic], [(lexical, 0.0)])

        self.assertEqual(len(ranked), 1)
        self.assertIs(ranked[0][0], lexical)

    def test_empty_rankings_produce_no_candidates(self) -> None:
        self.assertEqual(self.retriever._reciprocal_rank_fusion([], []), [])


class HybridRetrievalIntegrationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.semantic = make_document(
            "A vehicle is suitable for family transport.",
            "document:vehicle",
        )
        self.billing = make_document(
            "Invoice payment billing information.",
            "document:billing",
        )
        self.exact = make_document(
            "Reference ZX-9182 has been approved.",
            "document:exact",
        )
        self.unrelated = make_document(
            "Cloud storage backup configuration.",
            "document:cloud",
        )
        self.documents = [
            self.semantic,
            self.billing,
            self.exact,
            self.unrelated,
        ]

    def build_retriever(self, **settings: Any) -> HybridRetriever:
        vector_store = MetadataFilteredVectorStore(self.documents)
        return HybridRetriever(
            vector_store,
            lexical_documents=self.documents,
            **settings,
        )

    def test_real_vector_store_finds_a_semantic_synonym_without_keyword_overlap(
        self,
    ) -> None:
        retriever = self.build_retriever(vector_k=1, final_k=1)

        self.assertEqual(retriever._bm25_search("automobile"), [])
        self.assertEqual(document_sources(retriever.invoke("automobile")), ["document:vehicle"])

    def test_bm25_recovers_an_identifier_missing_from_top_vector_result(self) -> None:
        retriever = self.build_retriever(vector_k=1, final_k=3)

        vector_sources = document_sources(retriever._vector_search("invoice ZX-9182"))
        combined_sources = document_sources(retriever.invoke("invoice ZX-9182"))

        self.assertEqual(vector_sources, ["document:billing"])
        self.assertIn("document:exact", combined_sources)

    def test_shared_semantic_and_lexical_result_wins_rrf_before_reranking(self) -> None:
        retriever = self.build_retriever(vector_k=1, final_k=2)

        self.assertEqual(
            document_sources(retriever.invoke("invoice ZX-9182")),
            ["document:billing", "document:exact"],
        )

    def test_real_reranking_promotes_exact_identifier_before_final_cutoff(self) -> None:
        model = DeterministicCrossEncoder()
        reranker = CrossEncoderReranker(model_name="deterministic-cross-encoder")
        retriever = self.build_retriever(
            vector_k=1,
            final_k=1,
            rerank_k=2,
            reranker=reranker,
        )

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            results = retriever.invoke("invoice ZX-9182")

        self.assertEqual(document_sources(results), ["document:exact"])
        self.assertEqual(len(model.calls), 1)
        self.assertEqual(len(model.calls[0]["pairs"]), 2)

    def test_candidate_cutoff_prevents_reranking_unselected_documents(self) -> None:
        model = DeterministicCrossEncoder()
        retriever = self.build_retriever(
            vector_k=1,
            final_k=1,
            rerank_k=1,
            reranker=CrossEncoderReranker(),
        )

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            results = retriever.invoke("invoice ZX-9182")

        self.assertEqual(document_sources(results), ["document:billing"])
        self.assertEqual(model.calls, [])

    def test_source_filter_is_respected_by_real_vector_store_and_bm25(self) -> None:
        retriever = self.build_retriever(
            source="document:exact",
            vector_k=5,
            final_k=5,
        )

        results = retriever.invoke("invoice ZX-9182")

        self.assertEqual(document_sources(results), ["document:exact"])
        self.assertEqual(retriever.vector_store.last_search["filter"], {"source": "document:exact"})

    def test_same_chunk_returned_by_both_paths_appears_only_once(self) -> None:
        retriever = self.build_retriever(vector_k=3, final_k=6)

        results = retriever.invoke("invoice")
        identities = [
            (document.metadata["source"], document.metadata["chunk_index"])
            for document in results
        ]

        self.assertEqual(len(identities), len(set(identities)))

    def test_final_result_limit_is_applied_after_fusion(self) -> None:
        retriever = self.build_retriever(vector_k=4, final_k=2)

        self.assertEqual(len(retriever.invoke("invoice storage")), 2)

    def test_empty_vector_store_still_returns_real_bm25_matches(self) -> None:
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            lexical_documents=self.documents,
            final_k=1,
        )

        self.assertEqual(document_sources(retriever.invoke("ZX-9182")), ["document:exact"])

    def test_empty_retrieval_paths_return_an_empty_result(self) -> None:
        retriever = HybridRetriever(
            MetadataFilteredVectorStore([]),
            lexical_documents=[],
        )

        self.assertEqual(retriever.invoke("ZX-9182"), [])

    def test_reranker_failure_preserves_real_fused_result_order(self) -> None:
        baseline = self.build_retriever(vector_k=1, final_k=2)
        expected = document_sources(baseline.invoke("invoice ZX-9182"))
        model = DeterministicCrossEncoder(error=RuntimeError("inference unavailable"))
        retriever = self.build_retriever(
            vector_k=1,
            final_k=2,
            rerank_k=2,
            reranker=CrossEncoderReranker(),
        )

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            with self.assertLogs("rag.reranking", level="WARNING"):
                actual = document_sources(retriever.invoke("invoice ZX-9182"))

        self.assertEqual(actual, expected)


class CrossEncoderIntegrationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.documents = [
            make_document("General invoice payment details.", "document:generic"),
            make_document("The approved reference is ZX-9182.", "document:exact"),
            make_document("Other invoice information.", "document:other"),
        ]

    def test_actual_reranker_sorts_scores_from_predictor(self) -> None:
        model = DeterministicCrossEncoder()
        reranker = CrossEncoderReranker(batch_size=4)

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            ranked = reranker.rerank("invoice ZX-9182", self.documents)

        self.assertEqual(ranked[0].metadata["source"], "document:exact")
        self.assertEqual(model.calls[0]["batch_size"], 4)
        self.assertFalse(model.calls[0]["show_progress_bar"])

    def test_question_and_original_unicode_content_are_passed_without_mutation(
        self,
    ) -> None:
        question = "شماره قرارداد ۴۵۶ چیست؟"
        documents = [
            make_document("شماره قرارداد ٤٥٦ است.", "document:persian"),
            make_document("راهنمای پرداخت.", "document:payment"),
        ]
        model = DeterministicCrossEncoder()

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            CrossEncoderReranker().rerank(question, documents)

        self.assertEqual(
            model.calls[0]["pairs"],
            [(question, document.page_content) for document in documents],
        )

    def test_model_name_and_device_are_forwarded_to_model_loader(self) -> None:
        model = DeterministicCrossEncoder()
        reranker = CrossEncoderReranker(model_name="fixture-model", device="cpu")

        with patch("rag.reranking._load_cross_encoder", return_value=model) as load:
            reranker.rerank("invoice ZX-9182", self.documents)

        load.assert_called_once_with("fixture-model", "cpu")

    def test_lazy_model_loader_caches_model_instances_by_name_and_device(self) -> None:
        constructions: list[tuple[str, str | None]] = []

        class FixtureCrossEncoder:
            def __init__(self, model_name: str, *, device: str | None) -> None:
                constructions.append((model_name, device))

        module = ModuleType("sentence_transformers")
        module.CrossEncoder = FixtureCrossEncoder
        _load_cross_encoder.cache_clear()

        try:
            with patch.dict("sys.modules", {"sentence_transformers": module}):
                first = _load_cross_encoder("fixture-model", "cpu")
                repeated = _load_cross_encoder("fixture-model", "cpu")
                other = _load_cross_encoder("other-model", "cpu")

            self.assertIs(first, repeated)
            self.assertIsNot(first, other)
            self.assertEqual(
                constructions,
                [("fixture-model", "cpu"), ("other-model", "cpu")],
            )

        finally:
            _load_cross_encoder.cache_clear()

    def test_equal_model_scores_preserve_incoming_rrf_order(self) -> None:
        model = DeterministicCrossEncoder(scores=[0.5, 0.5, 0.5])

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            ranked = CrossEncoderReranker().rerank("invoice", self.documents)

        self.assertEqual(ranked, self.documents)

    def test_invalid_model_scores_keep_fused_candidates_intact(self) -> None:
        model = DeterministicCrossEncoder(scores=[0.9, "not-a-score", 0.1])

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            with self.assertLogs("rag.reranking", level="WARNING"):
                ranked = CrossEncoderReranker().rerank("invoice", self.documents)

        self.assertEqual(ranked, self.documents)

    def test_missing_model_scores_keep_fused_candidates_intact(self) -> None:
        model = DeterministicCrossEncoder(scores=[0.9])

        with patch("rag.reranking._load_cross_encoder", return_value=model):
            with self.assertLogs("rag.reranking", level="WARNING"):
                ranked = CrossEncoderReranker().rerank("invoice", self.documents)

        self.assertEqual(ranked, self.documents)


class StructuredDocumentRetrievalIntegrationTests(unittest.TestCase):
    @staticmethod
    def create_contract(path: Path) -> None:
        from docx import Document as DocxDocument

        document = DocxDocument()
        document.add_heading("شرایط قرارداد", level=1)
        document.add_heading("پرداخت", level=2)

        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "کد سفارش"
        table.cell(0, 1).text = "مبلغ"
        table.cell(1, 0).text = "ZX-9182"
        table.cell(1, 1).text = "۵۰ میلیون تومان"

        document.add_heading("امنیت", level=2)
        document.add_paragraph("رمز ورود باید محرمانه باقی بماند.")
        document.save(path)

    def test_docx_table_is_searchable_through_complete_hybrid_pipeline(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "contract.docx"
            self.create_contract(path)
            extracted = load_document(path)
            chunks = split_document_text(extracted, source="document:contract")
            retriever = HybridRetriever(
                MetadataFilteredVectorStore(chunks),
                lexical_documents=chunks,
                vector_k=3,
                final_k=1,
                rerank_k=3,
                reranker=CrossEncoderReranker(),
            )
            model = DeterministicCrossEncoder()

            with patch("rag.reranking._load_cross_encoder", return_value=model):
                results = retriever.invoke("مبلغ سفارش ZX-9182 چقدر است؟")

        self.assertEqual(len(results), 1)
        self.assertIn("۵۰ میلیون تومان", results[0].page_content)
        self.assertEqual(results[0].metadata["section_path"], "شرایط قرارداد > پرداخت")

    def test_serialized_reranked_source_keeps_original_section_and_citation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "contract.docx"
            self.create_contract(path)
            chunks = split_document_text(load_document(path), source="document:contract")
            retriever = HybridRetriever(
                MetadataFilteredVectorStore(chunks),
                lexical_documents=chunks,
                final_k=1,
            )

            serialized = serialize_sources(retriever.invoke("ZX-9182"))

        self.assertEqual(serialized[0]["source"], "document:contract")
        self.assertEqual(serialized[0]["section"], "شرایط قرارداد > پرداخت")
        self.assertIn("شرایط قرارداد > پرداخت", serialized[0]["citation"])


@unittest.skipUnless(
    os.getenv("RUN_REAL_RERANKER_TESTS", "").strip().casefold()
    in {"1", "true", "yes", "on"},
    "Set RUN_REAL_RERANKER_TESTS=1 to load the actual Cross-Encoder model.",
)
class RealCrossEncoderModelTests(unittest.TestCase):
    def test_downloaded_cross_encoder_promotes_the_relevant_answer(self) -> None:
        model_name = os.getenv("RERANKER_MODEL", DEFAULT_RERANKER_MODEL)
        device = os.getenv("RERANKER_DEVICE", "cpu")
        documents = [
            make_document(
                "A bicycle is repaired using a wheel pump and spare tires.",
                "document:unrelated",
            ),
            make_document(
                "Paris is the capital city of France.",
                "document:answer",
            ),
        ]

        model = _load_cross_encoder(model_name, device)
        self.assertTrue(callable(model.predict))

        ranked = CrossEncoderReranker(
            model_name=model_name,
            device=device,
            batch_size=2,
        ).rerank("What is the capital city of France?", documents)

        self.assertEqual(ranked[0].metadata["source"], "document:answer")
